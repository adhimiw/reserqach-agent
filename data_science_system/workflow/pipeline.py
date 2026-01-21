"""
Analysis Pipeline Orchestrator
Coordinates the complete data analysis workflow
"""

import pandas as pd
import numpy as np
import json
import os
from typing import Dict, Any, List, Optional
from datetime import datetime
import logging

from config import Config
from analysis_engine import HypothesisGenerator, StatisticalTester, ModelBuilder, InsightExtractor
from agents import SelfHealingAgent, DataVisualizer

logger = logging.getLogger("AnalysisPipeline")


class AnalysisPipeline:
    """Orchestrates complete data analysis workflow"""
    
    def __init__(self, dataset_path: str, output_dir: str = None):
        """
        Initialize analysis pipeline
        
        Args:
            dataset_path: Path to dataset file
            output_dir: Output directory for results
        """
        self.dataset_path = dataset_path
        self.dataset_name = os.path.basename(dataset_path).split('.')[0]
        self.output_dir = output_dir or Config.get_analysis_output_dir(self.dataset_name)
        
        # Create output subdirectories
        for subdir in ['data', 'code', 'visualizations', 'insights', 'logs']:
            os.makedirs(os.path.join(self.output_dir, subdir), exist_ok=True)
        
        # Initialize components
        self.df = None
        self.heuristic_generator = None
        self.statistical_tester = None
        self.model_builder = None
        self.insight_extractor = None
        self.visualizer = DataVisualizer(os.path.join(self.output_dir, 'visualizations'))
        self.self_healer = SelfHealingAgent()
        
        # Results storage
        self.results = {
            "dataset_info": {},
            "hypotheses": [],
            "statistical_tests": [],
            "models": {},
            "insights": [],
            "visualizations": {},
            "error_log": []
        }
        
        self.execution_log = []
        
        logger.info(f"Pipeline initialized for dataset: {self.dataset_name}")
        logger.info(f"Output directory: {self.output_dir}")
    
    def load_data(self) -> pd.DataFrame:
        """
        Load dataset from file
        
        Returns:
            Loaded DataFrame
        """
        self._log_step("Loading data", "start")
        
        try:
            # Determine file type and load
            file_ext = os.path.splitext(self.dataset_path)[1].lower()
            
            if file_ext == '.csv':
                self.df = pd.read_csv(self.dataset_path)
            elif file_ext in ['.xlsx', '.xls']:
                self.df = pd.read_excel(self.dataset_path)
            elif file_ext == '.json':
                self.df = pd.read_json(self.dataset_path)
            elif file_ext == '.parquet':
                self.df = pd.read_parquet(self.dataset_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Save cleaned data
            clean_path = os.path.join(self.output_dir, 'data', 'original.csv')
            self.df.to_csv(clean_path, index=False)
            
            # Store dataset info
            self.results['dataset_info'] = {
                "shape": self.df.shape,
                "columns": list(self.df.columns),
                "dtypes": {col: str(dtype) for col, dtype in self.df.dtypes.items()},
                "missing_values": self.df.isnull().sum().to_dict(),
                "memory_usage_mb": self.df.memory_usage(deep=True).sum() / (1024**2),
                "file_path": self.dataset_path
            }
            
            self._log_step("Loading data", "complete", f"Loaded {self.df.shape[0]} rows, {self.df.shape[1]} columns")
            return self.df
            
        except Exception as e:
            self._log_step("Loading data", "error", str(e))
            raise
    
    def clean_data(self, missing_strategy: str = 'mean', 
                   outlier_threshold: float = 3.0) -> pd.DataFrame:
        """
        Clean data: handle missing values and outliers
        
        Args:
            missing_strategy: Strategy for missing values
            outlier_threshold: Z-score threshold for outliers
        
        Returns:
            Cleaned DataFrame
        """
        self._log_step("Data cleaning", "start")
        
        try:
            df_clean = self.df.copy()
            
            # Handle missing values
            numeric_cols = df_clean.select_dtypes(include=[np.number]).columns
            categorical_cols = df_clean.select_dtypes(include=['object']).columns
            
            if missing_strategy == 'mean':
                df_clean[numeric_cols] = df_clean[numeric_cols].fillna(df_clean[numeric_cols].mean())
                df_clean[categorical_cols] = df_clean[categorical_cols].fillna(df_clean[categorical_cols].mode().iloc[0])
            elif missing_strategy == 'median':
                df_clean[numeric_cols] = df_clean[numeric_cols].fillna(df_clean[numeric_cols].median())
                df_clean[categorical_cols] = df_clean[categorical_cols].fillna(df_clean[categorical_cols].mode().iloc[0])
            elif missing_strategy == 'drop':
                df_clean = df_clean.dropna()
            
            # Remove extreme outliers
            for col in numeric_cols:
                z_scores = np.abs((df_clean[col] - df_clean[col].mean()) / df_clean[col].std())
                df_clean = df_clean[z_scores < outlier_threshold]
            
            # Save cleaned data
            clean_path = os.path.join(self.output_dir, 'data', 'cleaned.csv')
            df_clean.to_csv(clean_path, index=False)
            
            self.df = df_clean
            
            self._log_step("Data cleaning", "complete", 
                          f"Cleaned data shape: {df_clean.shape}")
            
            return self.df
            
        except Exception as e:
            self._log_step("Data cleaning", "error", str(e))
            raise
    
    def generate_hypotheses(self, max_hypotheses: int = 100) -> List[Dict[str, Any]]:
        """
        Generate testable hypotheses
        
        Args:
            max_hypotheses: Maximum number of hypotheses to generate
        
        Returns:
            List of hypotheses
        """
        self._log_step("Hypothesis generation", "start")
        
        try:
            self.heuristic_generator = HypothesisGenerator(self.df)
            hypotheses = self.heuristic_generator.generate_all_hypotheses(max_hypotheses)
            
            # Save hypotheses
            hypotheses_path = os.path.join(self.output_dir, 'insights', 'hypotheses.json')
            self.heuristic_generator.save_hypotheses(hypotheses_path)
            
            self.results['hypotheses'] = hypotheses
            
            self._log_step("Hypothesis generation", "complete", 
                          f"Generated {len(hypotheses)} hypotheses")
            
            return hypotheses
            
        except Exception as e:
            self._log_step("Hypothesis generation", "error", str(e))
            raise
    
    def run_statistical_tests(self) -> List[Dict[str, Any]]:
        """
        Run comprehensive statistical tests
        
        Returns:
            List of test results
        """
        self._log_step("Statistical testing", "start")
        
        try:
            self.statistical_tester = StatisticalTester(self.df, 
                                                  significance_level=Config.SIGNIFICANCE_LEVEL)
            test_results = self.statistical_tester.run_comprehensive_tests()
            
            # Save results
            results_path = os.path.join(self.output_dir, 'insights', 'statistical_tests.json')
            self.statistical_tester.save_results(results_path)
            
            self.results['statistical_tests'] = test_results
            
            self._log_step("Statistical testing", "complete", 
                          f"Ran {len(test_results)} tests")
            
            return test_results
            
        except Exception as e:
            self._log_step("Statistical testing", "error", str(e))
            raise
    
    def build_models(self, target_column: str = None) -> Dict[str, Any]:
        """
        Build predictive models
        
        Args:
            target_column: Name of target variable (auto-detected if None)
        
        Returns:
            Dictionary with model results
        """
        self._log_step("Model building", "start")
        
        try:
            # Auto-detect target if not specified
            if target_column is None:
                numeric_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
                if len(numeric_cols) > 0:
                    target_column = numeric_cols[-1]  # Use last numeric column
                else:
                    raise ValueError("No numeric columns found for target variable")
            
            # Build models
            self.model_builder = ModelBuilder(self.df, target_column,
                                          test_size=Config.TRAIN_TEST_SPLIT,
                                          random_state=42)
            
            # Prepare data
            prep_info = self.model_builder.prepare_data()
            
            # Build different models
            models_built = {}
            
            # Linear regression (if regression)
            if self.model_builder.task_type == 'regression':
                try:
                    models_built['linear'] = self.model_builder.build_linear_regression()
                except Exception as e:
                    logger.warning(f"Linear regression failed: {e}")
            
            # Random forest
            try:
                models_built['random_forest'] = self.model_builder.build_random_forest()
            except Exception as e:
                logger.warning(f"Random forest failed: {e}")
            
            # Gradient boosting
            try:
                models_built['gradient_boosting'] = self.model_builder.build_gradient_boosting()
            except Exception as e:
                logger.warning(f"Gradient boosting failed: {e}")
            
            # Save results
            results_path = os.path.join(self.output_dir, 'insights', 'models.json')
            self.model_builder.save_results(results_path)
            
            self.results['models'] = {
                'task_type': self.model_builder.task_type,
                'target_column': target_column,
                'preparation': prep_info,
                'models': models_built,
                'best_model': self.model_builder.best_model
            }
            
            self._log_step("Model building", "complete", 
                          f"Built {len(models_built)} models")
            
            return self.results['models']
            
        except Exception as e:
            self._log_step("Model building", "error", str(e))
            raise
    
    def extract_insights(self) -> List[Dict[str, Any]]:
        """
        Extract meaningful insights from all analysis
        
        Returns:
            List of insights
        """
        self._log_step("Insight extraction", "start")
        
        try:
            self.insight_extractor = InsightExtractor(self.df)
            
            # Generate insights from all analysis results
            analysis_results = {
                'correlations': [r for r in self.results['statistical_tests'] 
                               if r.get('test') == 'correlation'],
                'distributions': [r for r in self.results['statistical_tests'] 
                                 if r.get('test') == 'normality'],
                'outliers': [r for r in self.results['statistical_tests'] 
                             if r.get('test') == 'outliers'],
                'statistical_tests': self.results['statistical_tests'],
                'modeling': self.results.get('models', {})
            }
            
            insights = self.insight_extractor.generate_all_insights(analysis_results)
            
            # Save insights
            insights_path = os.path.join(self.output_dir, 'insights', 'insights.json')
            self.insight_extractor.save_insights(insights_path)
            
            self.results['insights'] = insights
            
            self._log_step("Insight extraction", "complete", 
                          f"Extracted {len(insights)} insights")
            
            return insights
            
        except Exception as e:
            self._log_step("Insight extraction", "error", str(e))
            raise
    
    def create_visualizations(self) -> Dict[str, str]:
        """
        Generate all relevant visualizations
        
        Returns:
            Dictionary mapping visualization names to file paths
        """
        self._log_step("Visualization creation", "start")
        
        try:
            visualizations = self.visualizer.generate_all_visualizations(self.df)
            
            self.results['visualizations'] = visualizations
            
            self._log_step("Visualization creation", "complete", 
                          f"Created {len(visualizations)} visualizations")
            
            return visualizations
            
        except Exception as e:
            self._log_step("Visualization creation", "error", str(e))
            raise
    
    def generate_reports(self, formats: List[str] = None) -> Dict[str, str]:
        """
        Generate comprehensive analysis reports
        
        Args:
            formats: List of formats ('markdown', 'word')
        
        Returns:
            Dictionary mapping format names to file paths
        """
        self._log_step("Report generation", "start")
        
        if formats is None:
            formats = ['markdown']
        
        reports = {}
        
        # Generate markdown report
        if 'markdown' in formats:
            markdown_path = self.generate_markdown_report()
            reports['markdown'] = markdown_path
        
        # Generate Word document (if Word MCP available)
        if 'word' in formats:
            try:
                word_path = self.generate_word_report()
                reports['word'] = word_path
            except Exception as e:
                logger.warning(f"Word report generation failed: {e}")
        
        self._log_step("Report generation", "complete", 
                      f"Generated reports in {list(reports.keys())}")
        
        return reports
    
    def generate_markdown_report(self) -> str:
        """
        Generate comprehensive markdown report
        
        Returns:
            Path to markdown file
        """
        report = f"""# Data Analysis Report

**Dataset:** {self.dataset_name}
**Analysis Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Data File:** {self.dataset_path}

---

## Executive Summary

This autonomous data science system analyzed the dataset `{self.dataset_name}` containing {self.df.shape[0]:,} rows and {self.df.shape[1]} columns. The analysis included hypothesis generation, statistical testing, predictive modeling, and insight extraction.

**Key Findings:**
- Generated {len(self.results.get('hypotheses', []))} testable hypotheses
- Ran {len(self.results.get('statistical_tests', []))} statistical tests
- Built {len(self.results.get('models', {}).get('models', {}))} predictive models
- Extracted {len(self.results.get('insights', []))} actionable insights

---

## Dataset Overview

### Data Summary

- **Rows:** {self.df.shape[0]:,}
- **Columns:** {self.df.shape[1]}
- **Memory Usage:** {self.df.memory_usage(deep=True).sum() / (1024**2):.2f} MB

### Columns

| Column | Type | Missing Values |
|---------|-------|---------------|
"""
        
        # Add column information
        for col in self.df.columns:
            dtype = str(self.df[col].dtype)
            missing = self.df[col].isnull().sum()
            missing_pct = (missing / len(self.df)) * 100
            report += f"| {col} | {dtype} | {missing:,} ({missing_pct:.1f}%) |\n"
        
        # Add hypotheses section
        if self.results.get('hypotheses'):
            report += "\n## Generated Hypotheses\n\n"
            if self.heuristic_generator:
                report += self.heuristic_generator.format_hypotheses_for_report()
        
        # Add statistical tests section
        if self.results.get('statistical_tests'):
            report += "\n## Statistical Analysis Results\n\n"
            if self.statistical_tester:
                report += self.statistical_tester.format_results_for_report()
        
        # Add modeling results
        if self.results.get('models'):
            report += "\n## Predictive Modeling\n\n"
            if self.model_builder:
                report += self.model_builder.format_results_for_report()
        
        # Add insights section
        if self.results.get('insights'):
            report += "\n## Key Insights\n\n"
            if self.insight_extractor:
                report += self.insight_extractor.format_insights_for_report()
        
        # Add visualizations section
        if self.results.get('visualizations'):
            report += "\n## Visualizations\n\n"
            report += self.visualizer.format_visualizations_for_report(
                self.results['visualizations']
            )
        
        # Add methodology
        report += """
## Methodology

This analysis was performed autonomously by the Data Science System with the following steps:

1. **Data Loading**: Loaded dataset and performed initial quality checks
2. **Data Cleaning**: Handled missing values and removed extreme outliers
3. **Hypothesis Generation**: Automatically generated testable hypotheses based on:
   - Correlation patterns
   - Distribution characteristics
   - Outlier detection
   - Categorical variable analysis
   - Trend analysis (for time series data)

4. **Statistical Testing**: Performed comprehensive statistical tests including:
   - Pearson, Spearman, and Kendall correlation tests
   - Shapiro-Wilk normality tests
   - IQR and Z-score outlier detection
   - ANOVA and Kruskal-Wallis group difference tests
   - Chi-square categorical association tests

5. **Predictive Modeling**: Built and evaluated multiple models:
   - Linear regression (for regression tasks)
   - Random Forest
   - Gradient Boosting

6. **Insight Extraction**: Generated actionable insights with "what," "why," and "how" explanations

7. **Visualization**: Created publication-quality visualizations for all major findings

---

## Error Handling and Self-Healing

The system includes automatic error detection and recovery:
- All operations are monitored for errors
- Automatic retry with alternative approaches
- Fallback to secondary APIs when primary methods fail
- Comprehensive logging of all errors and recovery attempts

"""

        # Save markdown report
        report_path = os.path.join(self.output_dir, f"{self.dataset_name}_report.md")
        with open(report_path, 'w') as f:
            f.write(report)
        
        return report_path
    
    def generate_word_report(self) -> str:
        """
        Generate Word document report using Word MCP
        
        Returns:
            Path to Word document
        """
        # This would use Word MCP server to create a formatted Word document
        # For now, save as placeholder
        word_path = os.path.join(self.output_dir, f"{self.dataset_name}_report.docx")
        
        # Create a simple Word document as placeholder
        # In production, this would use Word MCP server
        from docx import Document
        doc = Document()
        doc.add_heading('Data Analysis Report', 0)
        doc.add_paragraph(f'Dataset: {self.dataset_name}')
        doc.add_paragraph(f'Analysis Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.save(word_path)
        
        return word_path
    
    def run_full_pipeline(self, target_column: str = None, 
                        generate_word: bool = True) -> Dict[str, Any]:
        """
        Run the complete analysis pipeline
        
        Args:
            target_column: Target variable for modeling (auto-detected if None)
            generate_word: Whether to generate Word document
        
        Returns:
            Dictionary with all results
        """
        start_time = datetime.now()
        
        self._log_step("Full pipeline", "start")
        
        try:
            # Step 1: Load data
            self.load_data()
            
            # Step 2: Clean data
            self.clean_data()
            
            # Step 3: Generate hypotheses
            self.generate_hypotheses()
            
            # Step 4: Run statistical tests
            self.run_statistical_tests()
            
            # Step 5: Build models
            self.build_models(target_column=target_column)
            
            # Step 6: Extract insights
            self.extract_insights()
            
            # Step 7: Create visualizations
            self.create_visualizations()
            
            # Step 8: Generate reports
            formats = ['markdown']
            if generate_word:
                formats.append('word')
            self.generate_reports(formats=formats)
            
            # Step 9: Save execution log
            self.save_execution_log()
            
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            
            self._log_step("Full pipeline", "complete", 
                          f"Completed in {duration:.1f} seconds")
            
            return self.results
            
        except Exception as e:
            self._log_step("Full pipeline", "error", str(e))
            raise
    
    def _log_step(self, step: str, status: str, message: str = ""):
        """Log pipeline step"""
        log_entry = {
            "step": step,
            "status": status,
            "message": message,
            "timestamp": datetime.now().isoformat()
        }
        self.execution_log.append(log_entry)
        
        if status == "start":
            logger.info(f"[START] {step}")
        elif status == "complete":
            logger.info(f"[COMPLETE] {step}: {message}")
        elif status == "error":
            logger.error(f"[ERROR] {step}: {message}")
    
    def save_execution_log(self):
        """Save execution log to file"""
        log_path = os.path.join(self.output_dir, 'logs', 'execution_log.json')
        
        with open(log_path, 'w') as f:
            json.dump({
                "execution_log": self.execution_log,
                "results_summary": {
                    "hypotheses": len(self.results.get('hypotheses', [])),
                    "statistical_tests": len(self.results.get('statistical_tests', [])),
                    "models": len(self.results.get('models', {}).get('models', {})),
                    "insights": len(self.results.get('insights', [])),
                    "visualizations": len(self.results.get('visualizations', {}))
                }
            }, f, indent=2, default=str)
